import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency 'python-docx'. Please install with: pip install python-docx\n"
        f"Original error: {exc}"
    )


def set_document_fonts(document: Document) -> None:
    styles = document.styles
    normal_style = styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Microsoft YaHei"
    # Set East Asian font explicitly for better CN rendering
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal_font.size = Pt(11)


def add_cover(document: Document) -> None:
    section = document.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.8), Inches(0.8)
    section.left_margin, section.right_margin = Inches(1.0), Inches(1.0)

    # Logo
    logo_path_candidates = [
        os.path.join("public", "logo.png"),
        os.path.join("src", "assets", "logo.png"),
        os.path.join("src", "assets", "logo_inverted.png"),
    ]
    for logo_path in logo_path_candidates:
        if os.path.exists(logo_path):
            try:
                document.add_picture(logo_path, width=Inches(1.2))
            except Exception:
                pass
            break

    # Title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Cook | 投资人与用户白皮书")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(28)
    run.bold = True

    # Subtitle
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = p2.add_run("全栈双后端智能烹饪应用（含多模型RAG）")
    sub.font.name = "Microsoft YaHei"
    sub._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    sub.font.size = Pt(14)

    # Metadata
    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = p3.add_run(f"版本 v1.0  |  生成时间：{datetime.now().strftime('%Y-%m-%d')}  |  作者：AI-Cook 团队")
    meta.font.size = Pt(10)

    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    h = document.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if bold:
        run.bold = True


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_code_block(document: Document, code_text: str) -> None:
    for line in code_text.strip("\n").split("\n"):
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def build_document() -> Document:
    doc = Document()
    set_document_fonts(doc)
    add_cover(doc)

    # 目录
    add_heading(doc, "目录", level=1)
    add_bullets(
        doc,
        [
            "1. 项目概览",
            "2. 面向用户的核心功能",
            "3. 创新点（为何与众不同）",
            "4. 技术架构",
            "5. RAG 功能与使用",
            "6. 路线图与关键指标（面向投资人）",
            "7. 版权与致谢",
        ],
    )
    doc.add_page_break()

    # 1. 项目概览
    add_heading(doc, "1. 项目概览", level=1)
    add_paragraph(
        doc,
        (
            "AI-Cook 是一个集菜谱搜索与智能膳食指导为一体的全栈应用，"
            "提供统一的 Vue 3 前端及双后端（FastAPI AI 后端 + Flask 搜索后端）。"
            "应用专注于中文场景的精准检索与个性化 AI 对话咨询，并支持多模型 RAG 向量检索，"
            "帮助用户更快发现合适菜谱、制定健康饮食计划，亦为团队提供易部署、易扩展的工程化基础。"
        )
    )
    add_bullets(
        doc,
        [
            "统一界面：单页应用（SPA），响应式与动画增强的现代体验",
            "双后端：AI 对话/个性化 + 本地智能搜索，职责清晰、各司其职",
            "多模型 RAG：支持 Google、OpenAI、智谱等嵌入向量，FAISS 高效向量检索",
            "一键启动：firststart.bat / aicook.bat 简化安装与日常开发",
            "单一依赖源：根级 pyproject.toml 统一管理，稳定可靠",
        ],
    )

    # 2. 面向用户的核心功能
    add_heading(doc, "2. 面向用户的核心功能", level=1)
    add_paragraph(doc, "核心页面与功能概览：", bold=True)
    add_bullets(
        doc,
        [
            "首页：功能亮点、互动卡片与使用流程",
            "AI 营养师：流式对话、BMI/热量计算、会话历史、多模型切换、向量库状态、健康监测",
            "AI 推荐：根据身高体重、活动水平、偏好等生成个性化菜谱",
            "随机菜谱：平衡餐食生成、历史留存与营养汇总",
            "限定条件搜索：基于食材与厨具条件的智能匹配（模糊/严格/生存模式）",
            "美食展厅：语义搜索、搜索意图解析、瀑布流与高级过滤",
            "我的菜谱：收藏与历史管理、统计洞察、数据导出",
            "菜谱详情：配料/步骤/营养/相关推荐/设备清单与响应式排版",
        ],
    )

    # 3. 创新点
    add_heading(doc, "3. 创新点（为何与众不同）", level=1)
    add_bullets(
        doc,
        [
            "多模型 RAG：嵌入模型可插拔、动态切换，覆盖多语境与成本区间",
            "混合检索：BM25 + 模糊匹配 + 向量召回，多路召回与加权融合排序",
            "工程化可用性：单一依赖源、脚本化启动、结构化日志、可观测性增强",
            "中文体验优化：结巴分词、针对中文语料调优的嵌入模型选择",
            "前后端解耦：AI 能力与搜索能力独立演进，降低耦合提升吞吐",
        ],
    )

    # 4. 技术架构
    add_heading(doc, "4. 技术架构", level=1)
    add_paragraph(doc, "前端：Vue 3 + Vite + Pinia + Vue Router + Element Plus + Axios")
    add_paragraph(doc, "AI 后端（FastAPI）：FastAPI、Pydantic、SQLAlchemy、LangChain/Google GenAI、FAISS、Structlog")
    add_paragraph(doc, "搜索后端（Flask）：Flask、Sentence-Transformers、scikit-learn、rank-bm25、Jieba、FuzzyWuzzy")
    add_paragraph(doc, "数据与存储：SQLite（默认）/ PostgreSQL（生产）、本地 FAISS 索引、持久化会话/收藏")
    add_paragraph(doc, "运维与工具：Uvicorn/Gunicorn、Docker（可选）、一键脚本、统一依赖管理")
    add_paragraph(
        doc,
        "架构原则：关注点分离、可替换的 AI 提供商、异步高并发、可观测性与可移植性。",
    )

    # 5. RAG 功能与使用
    add_heading(doc, "5. RAG 功能与使用", level=1)
    add_paragraph(doc, "嵌入服务商：Google / OpenAI / 智谱（向量支持）；豆包、DeepSeek（对话支持）")
    add_paragraph(doc, "智能配置：环境变量、密钥校验、自动降级与性能监控")
    add_paragraph(doc, "前端操作：设置 → 选择服务商 → 填写密钥 → 一键生成 17,000+ 向量")
    add_paragraph(doc, "API 用法：")
    add_code_block(
        doc,
        """
GET /api/v1/rag/providers
POST /api/v1/rag/rebuild-embeddings?provider=doubao&batch_size=100
""",
    )
    add_paragraph(doc, "命令行用法：")
    add_code_block(
        doc,
        """
python aire-backend/scripts/import_data.py rebuild --embedding-provider zhipu
""",
    )

    # 6. 路线图与关键指标
    add_heading(doc, "6. 路线图与关键指标（面向投资人）", level=1)
    add_paragraph(doc, "建议关键指标（示例）：")
    add_bullets(
        doc,
        [
            "用户侧：DAU/WAU、留存率、搜索点击率、对话采纳率、转化（收藏/生成餐单）",
            "性能侧：端到端响应时间、RAG 向量重建耗时、检索召回率与相关性评分",
            "商业侧：订阅/增值服务渗透率、ARPU、API 成本占营收比",
        ],
    )
    add_paragraph(
        doc,
        (
            "结构化汇报建议参考 Bit.ai 对投资人更新的最佳实践（含‘创始人寄语/需要协助/成功与挑战/KPI’），"
            "并可结合 investory.io 建立指标与更新的看板化归档。"
        )
    )

    # 7. 版权与致谢
    add_heading(doc, "7. 版权与致谢", level=1)
    add_paragraph(
        doc,
        "本项目部分灵感与数据来自开源社区（致谢云游君‘做饭吧！’）。本白皮书基于仓库 README 与后端 RAG 说明整理。",
    )

    # 参考链接（外部资源）
    add_heading(doc, "附：参考链接", level=1)
    add_paragraph(
        doc,
        "投资人更新最佳实践（Bit.ai）：https://blog.bit.ai/investor-update-document/",
    )
    add_paragraph(
        doc,
        "投资人与指标管理（investory.io 帮助中心）：https://docs.investory.io/en/",
    )

    return doc


def main() -> None:
    output_dir = os.path.join("docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "AI-Cook_项目白皮书_投资人与用户版_v1.0.docx")

    doc = build_document()
    doc.save(output_path)
    print(f"Generated: {os.path.abspath(output_path)}")


if __name__ == "__main__":
    main()

